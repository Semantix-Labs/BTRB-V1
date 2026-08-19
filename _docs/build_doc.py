"""
Generates HANDOVER_DOCUMENT.docx from the markdown source.
Run: python3 _docs/build_doc.py
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import re, copy

# ── Brand colours ──────────────────────────────────────────────────────────────
NAVY      = RGBColor(0x1A, 0x3A, 0x61)   # primary
ACCENT    = RGBColor(0x2A, 0x7A, 0xE4)   # accent blue
TABLE_HDR = RGBColor(0x1A, 0x3A, 0x61)   # table header bg
TABLE_ALT = RGBColor(0xF3, 0xF6, 0xFB)   # alternate row
CODE_BG   = RGBColor(0xF1, 0xF5, 0xF9)
WARN_BG   = RGBColor(0xFF, 0xF7, 0xED)
WARN_BRD  = RGBColor(0xFB, 0x92, 0x3C)
NOTE_BG   = RGBColor(0xEF, 0xF6, 0xFF)
NOTE_BRD  = RGBColor(0x60, 0xA5, 0xFA)
WHITE     = RGBColor(0xFF, 0xFF, 0xFF)
GRAY_TXT  = RGBColor(0x6B, 0x72, 0x80)
DARK_TXT  = RGBColor(0x11, 0x18, 0x27)

# ── Helpers ────────────────────────────────────────────────────────────────────

def rgb_hex(color) -> str:
    return f'{color[0]:02X}{color[1]:02X}{color[2]:02X}'

def set_cell_bg(cell, color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), rgb_hex(color))
    tcPr.append(shd)

def set_cell_border(cell, top=None, bottom=None, left=None, right=None):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for side, color in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        if color:
            el = OxmlElement(f'w:{side}')
            el.set(qn('w:val'), 'single')
            el.set(qn('w:sz'), '4')
            el.set(qn('w:space'), '0')
            el.set(qn('w:color'), rgb_hex(color))
            tcBorders.append(el)
    tcPr.append(tcBorders)

def paragraph_border_left(para, color, width_pt=18):
    pPr = para._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    left = OxmlElement('w:left')
    left.set(qn('w:val'), 'single')
    left.set(qn('w:sz'), str(width_pt))
    left.set(qn('w:space'), '6')
    left.set(qn('w:color'), rgb_hex(color))
    pBdr.append(left)
    pPr.append(pBdr)

def shade_paragraph(para, color):
    pPr = para._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), rgb_hex(color))
    pPr.append(shd)

def add_run_bold(para, text, color=None, size=None):
    run = para.add_run(text)
    run.bold = True
    if color:
        run.font.color.rgb = color
    if size:
        run.font.size = Pt(size)
    return run

def set_para_spacing(para, before=0, after=4, line=None):
    pPr = para._p.get_or_add_pPr()
    spacing = OxmlElement('w:spacing')
    spacing.set(qn('w:before'), str(before * 20))
    spacing.set(qn('w:after'), str(after * 20))
    if line:
        spacing.set(qn('w:line'), str(int(line * 240)))
        spacing.set(qn('w:lineRule'), 'auto')
    pPr.append(spacing)

def apply_inline_formatting(para, text, base_color=DARK_TXT, base_size=10.5):
    """Parse **bold**, `code`, and plain text segments and add runs."""
    # pattern order: bold, code, plain
    pattern = re.compile(r'(\*\*(.+?)\*\*|`([^`]+)`)')
    pos = 0
    for m in pattern.finditer(text):
        # plain text before match
        if m.start() > pos:
            r = para.add_run(text[pos:m.start()])
            r.font.size = Pt(base_size)
            r.font.color.rgb = base_color
        if m.group(0).startswith('**'):
            r = para.add_run(m.group(2))
            r.bold = True
            r.font.size = Pt(base_size)
            r.font.color.rgb = base_color
        else:  # backtick code
            r = para.add_run(m.group(3))
            r.font.name = 'Courier New'
            r.font.size = Pt(9)
            r.font.color.rgb = RGBColor(0xBE, 0x18, 0x5D)
        pos = m.end()
    # trailing plain text
    if pos < len(text):
        r = para.add_run(text[pos:])
        r.font.size = Pt(base_size)
        r.font.color.rgb = base_color

# ── Document setup ─────────────────────────────────────────────────────────────

doc = Document()

# Page margins
for section in doc.sections:
    section.top_margin    = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin   = Cm(2.5)
    section.right_margin  = Cm(2.5)

# Default style
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(10.5)
style.font.color.rgb = DARK_TXT

# ── Cover page ─────────────────────────────────────────────────────────────────

def add_cover(doc):
    # Top colour bar  ─ via a 1-cell table
    tbl = doc.add_table(rows=1, cols=1)
    tbl.style = 'Table Grid'
    cell = tbl.rows[0].cells[0]
    set_cell_bg(cell, NAVY)
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(28)
    p.paragraph_format.space_after  = Pt(28)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run('BARB SL WEBSITE')
    r.font.name  = 'Calibri'
    r.font.bold  = True
    r.font.size  = Pt(28)
    r.font.color.rgb = WHITE
    r2 = p.add_run('\nProject Handover Document')
    r2.font.name  = 'Calibri'
    r2.font.size  = Pt(16)
    r2.font.color.rgb = RGBColor(0xBF, 0xD7, 0xFF)
    # remove table border
    for row in tbl.rows:
        for c in row.cells:
            for side in ('top','bottom','left','right'):
                el = OxmlElement(f'w:{side}')
                tc = c._tc; tcPr = tc.get_or_add_tcPr()
                tcBorders = OxmlElement('w:tcBorders')
                el.set(qn('w:val'), 'none')
                tcBorders.append(el)
                tcPr.append(tcBorders)

    doc.add_paragraph()

    meta = [
        ('Prepared by',      'Semantix Labs'),
        ('Prepared for',     'Behaviour Analyst Registration Board Sri Lanka (BARB SL)'),
        ('Project',          'BTRB-V1 — Official BARB SL Website'),
        ('Date',             'July 2026'),
        ('Document version', '1.1'),
    ]
    for label, value in meta:
        p = doc.add_paragraph()
        r1 = p.add_run(f'{label}:  ')
        r1.bold = True
        r1.font.color.rgb = NAVY
        r1.font.size = Pt(10.5)
        r2 = p.add_run(value)
        r2.font.size = Pt(10.5)
        set_para_spacing(p, before=2, after=2)

    doc.add_page_break()

add_cover(doc)

# ── Section heading helpers ────────────────────────────────────────────────────

def add_h1(doc, text):
    p = doc.add_paragraph()
    set_para_spacing(p, before=18, after=6)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(18)
    run.font.color.rgb = NAVY
    run.font.name = 'Calibri'
    # bottom border
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bot = OxmlElement('w:bottom')
    bot.set(qn('w:val'), 'single')
    bot.set(qn('w:sz'), '6')
    bot.set(qn('w:space'), '1')
    bot.set(qn('w:color'), rgb_hex(NAVY))
    pBdr.append(bot)
    pPr.append(pBdr)
    return p

def add_h2(doc, text):
    p = doc.add_paragraph()
    set_para_spacing(p, before=14, after=4)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(13.5)
    run.font.color.rgb = NAVY
    run.font.name = 'Calibri'
    return p

def add_h3(doc, text):
    p = doc.add_paragraph()
    set_para_spacing(p, before=10, after=3)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(11.5)
    run.font.color.rgb = ACCENT
    run.font.name = 'Calibri'
    return p

def add_h4(doc, text):
    p = doc.add_paragraph()
    set_para_spacing(p, before=8, after=2)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(10.5)
    run.font.color.rgb = NAVY
    run.font.name = 'Calibri'
    return p

def add_body(doc, text, indent=False):
    p = doc.add_paragraph()
    set_para_spacing(p, before=0, after=4, line=1.15)
    if indent:
        p.paragraph_format.left_indent = Inches(0.25)
    apply_inline_formatting(p, text)
    return p

def add_bullet(doc, text, level=0):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.left_indent  = Inches(0.25 + level * 0.2)
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after  = Pt(1)
    apply_inline_formatting(p, text)
    return p

def add_numbered(doc, text, num):
    p = doc.add_paragraph()
    set_para_spacing(p, before=1, after=1)
    p.paragraph_format.left_indent   = Inches(0.3)
    p.paragraph_format.first_line_indent = Inches(-0.3)
    r0 = p.add_run(f'{num}.  ')
    r0.bold = True
    r0.font.color.rgb = NAVY
    r0.font.size = Pt(10.5)
    apply_inline_formatting(p, text)
    return p

def add_code_block(doc, lines):
    for line in lines:
        p = doc.add_paragraph()
        shade_paragraph(p, CODE_BG)
        p.paragraph_format.left_indent = Inches(0.2)
        set_para_spacing(p, before=1, after=1)
        r = p.add_run(line if line else ' ')
        r.font.name = 'Courier New'
        r.font.size = Pt(9)
        r.font.color.rgb = RGBColor(0x1E, 0x29, 0x3B)

def add_note_box(doc, text, kind='note'):
    """kind = 'note' (blue) or 'warn' (orange)"""
    bg  = NOTE_BG   if kind == 'note' else WARN_BG
    brd = NOTE_BRD  if kind == 'note' else WARN_BRD
    p = doc.add_paragraph()
    shade_paragraph(p, bg)
    paragraph_border_left(p, brd, width_pt=18)
    p.paragraph_format.left_indent = Inches(0.15)
    set_para_spacing(p, before=4, after=4, line=1.15)
    apply_inline_formatting(p, text.lstrip('> ').strip())
    return p

def add_table(doc, headers, rows, col_widths=None):
    n_cols = len(headers)
    tbl = doc.add_table(rows=1 + len(rows), cols=n_cols)
    tbl.style = 'Table Grid'

    # Header row
    hdr_cells = tbl.rows[0].cells
    for i, h in enumerate(headers):
        set_cell_bg(hdr_cells[i], TABLE_HDR)
        p = hdr_cells[i].paragraphs[0]
        r = p.add_run(h)
        r.bold = True
        r.font.color.rgb = WHITE
        r.font.size = Pt(9.5)
        r.font.name = 'Calibri'
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after  = Pt(4)

    # Data rows
    for ri, row_data in enumerate(rows):
        cells = tbl.rows[ri + 1].cells
        bg = TABLE_ALT if ri % 2 == 1 else WHITE
        for ci, cell_text in enumerate(row_data):
            set_cell_bg(cells[ci], bg)
            p = cells[ci].paragraphs[0]
            apply_inline_formatting(p, str(cell_text), base_size=9.5)
            p.paragraph_format.space_before = Pt(3)
            p.paragraph_format.space_after  = Pt(3)

    # Column widths
    if col_widths:
        for row in tbl.rows:
            for i, w in enumerate(col_widths):
                row.cells[i].width = Inches(w)

    doc.add_paragraph()  # spacing after table
    return tbl

def add_divider(doc):
    p = doc.add_paragraph()
    set_para_spacing(p, before=6, after=6)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bot = OxmlElement('w:bottom')
    bot.set(qn('w:val'), 'single')
    bot.set(qn('w:sz'), '4')
    bot.set(qn('w:space'), '1')
    bot.set(qn('w:color'), 'CCCCCC')
    pBdr.append(bot)
    pPr.append(pBdr)

# ── TOC page ───────────────────────────────────────────────────────────────────

add_h1(doc, 'Table of Contents')
toc_items = [
    ('1', 'Project Overview'),
    ('2', 'Technology Stack'),
    ('3', 'Site Structure — All Pages'),
    ('4', 'How the Contact Form Works'),
    ('5', 'How the Certification Application Form Works'),
    ('6', 'How the Admin Panel Works'),
    ('7', 'Database Architecture'),
    ('8', 'File Storage'),
    ('9', 'Environment Variables & Configuration'),
    ('10', 'Deployment'),
    ('11', 'Credentials & Access Handover'),
    ('12', 'Ongoing Maintenance Notes'),
]
for num, title in toc_items:
    p = doc.add_paragraph()
    set_para_spacing(p, before=2, after=2)
    r1 = p.add_run(f'{num}.  ')
    r1.bold = True
    r1.font.color.rgb = NAVY
    r1.font.size = Pt(10.5)
    r2 = p.add_run(title)
    r2.font.size = Pt(10.5)
doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 1 — Project Overview
# ══════════════════════════════════════════════════════════════════════════════

add_h1(doc, '1. Project Overview')
add_body(doc, 'The BARB SL website is the official digital presence of the **Behaviour Analyst Registration Board Sri Lanka**. It serves as:')
for item in [
    'A **public information portal** about behaviour therapy in Sri Lanka',
    'A **certification gateway** for therapists to apply for BARB registration',
    'A **therapist directory** for the public to find BARB-registered practitioners',
    'A **fundraising & donation platform** for the organisation',
    'A **contact & inquiry hub** for all stakeholder communication',
    'An **administrative back office** for the BARB team to review applications, manage the therapist registry, and respond to inquiries',
]:
    add_bullet(doc, item)
add_body(doc, 'The website was designed and built from scratch by Semantix Labs in 2025–2026 based on BARB\'s brand guidelines, information architecture documentation, and content provided by the BARB team.')
add_divider(doc)

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 2 — Technology Stack
# ══════════════════════════════════════════════════════════════════════════════

add_h1(doc, '2. Technology Stack')
add_table(doc,
    headers=['Layer', 'Technology', 'Version'],
    rows=[
        ['Framework',        'Next.js (App Router)',        '16.1.2'],
        ['Language',         'TypeScript',                  '5.x'],
        ['Styling',          'Tailwind CSS',                '4.x'],
        ['UI Components',    'Radix UI + Lucide Icons',     'Latest'],
        ['Database & Auth',  'Supabase (PostgreSQL)',       'Latest'],
        ['File Storage',     'Supabase Storage',            '—'],
        ['Email Service',    'Resend',                      '—'],
        ['Containerisation', 'Docker',                      '—'],
        ['Runtime',          'Node.js',                     '20 (Alpine)'],
    ],
    col_widths=[1.6, 2.6, 1.4]
)
add_body(doc, '**Why these choices?**')
for item in [
    '**Next.js** gives the site fast server-side rendering, SEO-friendly pages, and built-in API routes — all in one project.',
    '**Supabase** provides a managed PostgreSQL database, authentication, file storage, and Row Level Security (RLS) policies, removing the need for a separate server.',
    '**Resend** handles all transactional emails (OTP codes for applicants, inquiry notifications for admins) reliably.',
    '**Docker** makes deployment consistent and environment-independent.',
]:
    add_bullet(doc, item)
add_divider(doc)

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 3 — Site Structure
# ══════════════════════════════════════════════════════════════════════════════

add_h1(doc, '3. Site Structure — All Pages')
add_body(doc, 'The website is organised into the following public-facing pages and one protected admin area.')

add_h2(doc, 'Public Pages')
add_table(doc,
    headers=['URL Path', 'Page Name', 'Purpose'],
    rows=[
        ['/',                       'Home',                'Landing page. Hero, stats, intro, therapy overview, ethics section, mailing list CTA.'],
        ['/about',                  'About Us',            'Mission, vision, problem statement, values, how we work, board members, governance, ethics committee, final CTA.'],
        ['/certification',          'Certification',       'Full certification pathway — eligibility, criteria, process steps, grandparenting pathway, experienced practitioners, previously certified, FAQ. Includes the "Apply Now" button.'],
        ['/certification/apply',    'Application Form',    'The multi-step certification application form (see Section 5).'],
        ['/training',               'Training',            'Training programs offered, upcoming programs, training benefits, value proposition, collaboration section, partners, CTA.'],
        ['/therapy',                'Behaviour Therapy',   'Educational content about what behaviour therapy is, who it helps, and why it matters.'],
        ['/directory',              'Therapist Directory', 'Searchable, filterable public registry of all BARB-registered therapists.'],
        ['/resources',              'Resources',           'General resource library for the public and professionals.'],
        ['/donate',                 'Donate',              'Donation tiers, ways to support, impact snapshot, sponsorship CTA, mission context, bank transfer details.'],
        ['/contact',                'Contact',             'Contact form + contact info panel. Covers Certification, Training, Donation, Collaboration, Complaints, General, and Other inquiry types.'],
        ['/legal/terms',            'Terms & Conditions',  'Legal terms of use.'],
        ['/legal/ethical-standards','Ethical Standards',   "BARB's published ethical standards document."],
    ],
    col_widths=[1.9, 1.5, 3.2]
)

add_h2(doc, 'Admin Pages (Password Protected)')
add_table(doc,
    headers=['URL Path', 'Page Name', 'Purpose'],
    rows=[
        ['/admin/login',                'Admin Login',           'Secure login page for BARB administrators.'],
        ['/admin/reset-password',       'Reset Password',        'Password reset page reached via a Supabase recovery email link.'],
        ['/admin',                      'Applications Dashboard','Lists all certification applications with status, search, and filter.'],
        ['/admin/applications/[id]',    'Application Detail',    'Full view of a single applicant\'s information with Approve / Reject action buttons.'],
        ['/admin/therapists',           'Therapist Registry',    'Lists all approved therapists with status management and delete capability.'],
        ['/admin/inquiries',            'Inquiries Dashboard',   'Lists all contact form submissions with search, filter, status management, and inline reply.'],
    ],
    col_widths=[1.9, 1.5, 3.2]
)
add_divider(doc)

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 4 — Contact Form
# ══════════════════════════════════════════════════════════════════════════════

add_h1(doc, '4. How the Contact Form Works')
add_body(doc, 'The contact form is located at `/contact`. It allows any visitor to send a message to BARB.')

add_h2(doc, 'What the user fills in')
add_table(doc,
    headers=['Field', 'Required', 'Notes'],
    rows=[
        ['Full Name',                 'Yes',         '—'],
        ['Email',                     'Yes',         '—'],
        ['Phone Number',              'Yes',         '—'],
        ['Purpose of Contact',        'Yes',         'Dropdown: Certification, Training, Donation, Collaboration, Complaints, General, Other'],
        ['Message',                   'Conditional', 'Required for all topics except Donation'],
        ['Privacy consent checkbox',  'Yes',         'Must be checked before submission'],
    ],
    col_widths=[1.8, 1.0, 3.8]
)
add_body(doc, '**Special behaviour for Donation topic:** When "Donation" is selected, an additional section appears dynamically asking the user to choose a donation tier (First Tier: $500, Second Tier: $300, Third Tier: $100, or Other custom amount). This tier information is prepended to the message when stored in the database.')

add_h2(doc, 'What happens when the form is submitted')
for i, step in enumerate([
    'The browser sends the form data to the API endpoint `/api/contact` via a POST request.',
    'The API validates that `name`, `email`, and `inquiry_type` are present.',
    'The inquiry is inserted into the `inquiries` table in Supabase with status `new`.',
    'An automated **email notification** is sent to the BARB admin email address (configured via `ADMIN_NOTIFICATION_EMAIL`) using the **Resend** email service. The email contains all form details and a direct link to `/admin/inquiries`.',
    'The user sees a green success screen: "Thank you for reaching out. We aim to respond within 3–5 business days."',
], 1):
    add_numbered(doc, step, i)

add_h2(doc, 'If the submission fails')
add_body(doc, 'The user sees a red error message on the form. Common causes: missing required fields, network issue, or Supabase/Resend service outage.')
add_divider(doc)

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 5 — Application Form
# ══════════════════════════════════════════════════════════════════════════════

add_h1(doc, '5. How the Certification Application Form Works')
add_body(doc, 'The application form is located at `/certification/apply` and is launched via the "Apply Now" button on the Certification page. It is a **5-step wizard** that collects all information needed for BARB registration.')

add_h2(doc, 'Step-by-step breakdown')

add_h3(doc, 'Step 1 — Personal Information')
add_body(doc, 'Collects the applicant\'s:')
for item in [
    'First name, Surname',
    'Date of Birth',
    'Address (line, city, post code)',
    'Primary and optional secondary phone number',
    'Email address',
    'NIC or Passport number',
    'Upload: NIC front and NIC back (or passport)',
]:
    add_bullet(doc, item)
add_body(doc, '**Email Verification (OTP):** After the user fills in Step 1 and clicks "Next", the system sends a **6-digit one-time password (OTP)** to the applicant\'s email address via Resend. A modal popup appears asking the user to enter the code. The code expires in **10 minutes**. Once verified, the form advances to Step 2. This prevents fake or mistyped email addresses.')

add_h3(doc, 'Step 2 — Criteria Selection')
add_body(doc, 'The applicant selects which criteria they are applying under. Options include:')
for item in [
    'Current RBT (Registered Behaviour Technician) — with certification number',
    'Current IBT (International Behaviour Technician) — with certification number',
    'Expired RBT — with document upload',
    'Voluntary Inactive RBT — with certification number and reactivation date',
    'Expired IBT — with document upload',
    'Practicing Behaviour Therapist',
    'Other ABA Qualifications',
    'Behaviour Analyst — with document upload',
]:
    add_bullet(doc, item)

add_h3(doc, 'Step 3 — Additional Information')
add_body(doc, 'Collects education and work experience details:')
for item in [
    '**Education:** Institution name, period of study, qualifications/degree, education certificate upload',
    '**Work Experience:** Workplace name, address, employment period, designation, full-time/part-time status, explanation of services, work experience letter upload',
    '**Other Documents:** CV upload, professional liability/malpractice insurance upload',
]:
    add_bullet(doc, item)

add_h3(doc, 'Step 4 — Terms & Conditions')
add_body(doc, 'The applicant must tick all mandatory agreement checkboxes:')
for item in [
    'Confirms residency eligibility',
    'Agrees to BARB\'s objectives',
    'Agrees to maintain standards',
    'Agrees to licensing terms',
    'Agrees to keep information updated',
    'Agrees to malpractice standards',
    'Agrees to BARB\'s ethical code',
    'Confirms police clearance compliance',
]:
    add_bullet(doc, item)

add_h3(doc, 'Step 5 — Review & Submit')
add_body(doc, 'Shows a summary of the applicant\'s information (name, email, phone, address, specialisation, number of documents attached). The applicant clicks **"Submit Application"**.')
add_body(doc, '**What happens on submission:**')
for i, step in enumerate([
    'All uploaded files are uploaded to Supabase Storage bucket `application-documents` in organised folders (`nic/`, `certifications/`, `education/`, `work/`, `cv/`, `insurance/`).',
    'A new record is inserted into the `therapist_applications` table with `review_status = \'pending\'`.',
    'The applicant sees a green confirmation screen: "Application Submitted! Your application has been received and is under review. We will contact you at [their email]."',
], 1):
    add_numbered(doc, step, i)
add_divider(doc)

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 6 — Admin Panel
# ══════════════════════════════════════════════════════════════════════════════

add_h1(doc, '6. How the Admin Panel Works')
add_body(doc, 'The admin panel is accessible at `/admin`. It is **password-protected** and only accessible to BARB staff with a Supabase Auth account.')

add_h2(doc, '6.1  Navigation Bar')
add_body(doc, 'The top navigation bar is shared across all admin pages. It has three sections:')
for item in [
    '**Left:** BARB Admin branding with a shield icon',
    '**Centre:** Navigation tabs — **Applications**, **Therapists**, **Inquiries** — displayed as pill-shaped buttons. The active tab is highlighted. The Inquiries tab shows a red badge with the count of unread inquiries.',
    '**Right:** The logged-in admin\'s email address (truncated) and a Logout button',
]:
    add_bullet(doc, item)

add_h2(doc, '6.2  Logging In — /admin/login')
add_body(doc, 'Enter the admin **email** and **password** registered in Supabase Authentication. On success, the user is redirected to the Applications Dashboard. The session persists across page refreshes.')
add_body(doc, 'To create admin accounts: go to Supabase Dashboard → Authentication → Users → "Invite user" or "Add user".')

add_h2(doc, '6.3  Resetting a Password — /admin/reset-password')
add_body(doc, 'This page is reached by clicking a password recovery link sent to the admin\'s email from Supabase.')
for item in [
    'While the recovery token is being verified, the page shows a spinner ("Verifying reset link…")',
    'Once verified, the admin enters a new password (minimum 8 characters) and a confirmation',
    'A show/hide password toggle is available',
    'On success, a green confirmation screen appears and the page automatically redirects to `/admin/login` after 3 seconds',
    'If the link has already expired, the page prompts the admin to request a new one',
]:
    add_bullet(doc, item)

add_h2(doc, '6.4  Applications Dashboard — /admin')
add_body(doc, 'This page shows all certification applications submitted via the website.')
add_h3(doc, 'Features')
for item in [
    '**Filter tabs:** All / Pending / In Progress / Approved / Not Approved — click to filter the list instantly. Each tab shows its count.',
    '**Search bar:** Type a name, email, or city to filter in real time.',
    '**Status badge / dropdown:** Pending and In Progress are editable via a dropdown. Approved and Not Approved are **locked** — they display as static badges and cannot be changed.',
    '**View button:** Opens the full Application Detail page.',
]:
    add_bullet(doc, item)
add_h3(doc, 'Status workflow')
add_code_block(doc, [
    'Pending  →  In Progress  →  Approved',
    '                         →  Not Approved',
])
add_table(doc,
    headers=['Status', 'Colour', 'Meaning'],
    rows=[
        ['Pending',      'Yellow', 'Just received, not yet actioned'],
        ['In Progress',  'Blue',   'Admin is actively reviewing this application'],
        ['Approved',     'Green',  'Application approved; therapist added to registry'],
        ['Not Approved', 'Red',    'Application declined'],
    ],
    col_widths=[1.3, 1.0, 4.3]
)
add_note_box(doc, '**Important:** Once an application is moved to Approved or Not Approved, the status is permanently locked. The decision is final. Use the Therapist Registry page to manage an approved therapist\'s ongoing status.', kind='warn')

add_h2(doc, '6.5  Application Detail Page — /admin/applications/[id]')
add_body(doc, 'Opens the full record for a single applicant. Accessible by clicking "View" on any application row.')
add_h3(doc, 'Displays')
for item in [
    '**Personal Information:** Date of birth, NIC/Passport number, phone numbers',
    '**Address:** Street address, city, post code',
    '**Certification Criteria:** Colour-coded badges for each criterion selected (e.g. "Current RBT", "Behaviour Analyst", "Expired IBT")',
    '**RBT / IBT Certification Numbers** (if provided)',
    '**Education:** Institution, period, qualifications',
    '**Work Experience:** Workplace, address, designation, employment period, full/part-time, explanation of services',
]:
    add_bullet(doc, item)

add_h3(doc, 'Decision Panel')
add_body(doc, 'Visible in the right sidebar while the status is Pending or In Progress:')
for item in [
    '**"Approve & Add to Directory"** — Marks the application as `approved`, creates a new therapist record with a generated registration number, and makes the therapist visible in the public directory. A green banner shows the assigned registration number.',
    '**"Reject Application"** — Marks the application as `rejected`. The applicant does not appear in the public directory.',
    'Once either action is taken, the buttons are hidden and replaced with a note showing the final decision.',
]:
    add_bullet(doc, item)

add_h3(doc, 'Registration Number Generation Logic')
add_body(doc, 'Registration numbers are generated automatically on approval. The format is `PREFIX-NNN` (e.g. `RBT-007`).')
add_h4(doc, 'Step 1 — Determine the prefix from the criteria checkboxes')
add_table(doc,
    headers=['Criteria checked', 'Prefix assigned'],
    rows=[
        ['`behaviour_analyst`',                                           '`RBA`'],
        ['`current_ibt` or `expired_ibt`',                               '`IBT`'],
        ['Everything else (RBT, practicing therapist, voluntary inactive, etc.)', '`RBT`'],
    ],
    col_widths=[3.5, 1.5]
)
add_h4(doc, 'Step 2 — Find the next sequential number for that prefix')
for i, step in enumerate([
    'Query the `therapists` table for all rows where `registration_number LIKE \'RBT-%\'` (substituting the actual prefix)',
    'Strip the prefix from each result: `\'RBT-007\'` → `\'007\'` → parseInt → `7`',
    'Take the maximum: e.g. `7`',
    'Add 1 and zero-pad to 3 digits: `8` → `RBT-008`',
    'First-ever approval for a prefix (no existing rows) starts at `001`',
], 1):
    add_numbered(doc, step, i)
add_body(doc, '**Prefix sequences are independent** — `RBT`, `IBT`, and `RBA` each have their own counter starting from `001`.')
add_note_box(doc, '**Compatibility with old records:** The database contains therapists with the legacy format (`BARB-2026-NNN`). The `LIKE \'RBT-%\'` filter does not match them, so the first approval after this change correctly starts at `RBT-001`.', kind='note')
add_note_box(doc, '**Re-approval edge case:** If a therapist record already exists for the application, the system skips number generation entirely and simply makes the existing record visible again. No duplicate number is created.', kind='note')
add_note_box(doc, '**Concurrency note:** If two applications are approved at the same time, both could read the same max and generate a duplicate number. In practice this is extremely unlikely. If guaranteed uniqueness is required, add database sequences in Supabase:\n\nCREATE SEQUENCE rbt_seq START 1;\nCREATE SEQUENCE rba_seq START 1;\nCREATE SEQUENCE ibt_seq START 1;\n\nThis is not currently implemented.', kind='warn')
add_body(doc, '**Contact Applicant section:** Direct mailto: link pre-filled with the applicant\'s email, and their phone number.')

add_h2(doc, '6.6  Therapist Registry — /admin/therapists')
add_body(doc, 'A management page for the full list of approved therapists, reached via the "Therapists" navigation tab.')
add_body(doc, '**Purpose:** Manage the ongoing status of therapists who have already been approved. Unlike the Applications page (which handles new submissions), this page is for managing existing registry members.')
add_h3(doc, 'Features')
for item in [
    '**Summary bar:** Shows total therapists and a count breakdown by status (Active / Inactive / Approved Professional)',
    '**Filter tabs:** All / Active / Inactive / Approved Professional — with counts',
    '**Search bar:** Search by name, email, registration number, or city',
    '**Table columns:** Therapist name + email, Registration Number, City, Designation, Status, Actions',
]:
    add_bullet(doc, item)
add_h3(doc, 'Status management (inline dropdown)')
add_table(doc,
    headers=['Status', 'Colour', 'Directory visibility', 'Meaning'],
    rows=[
        ['Active',               'Blue',  'Visible', 'Fully registered and practising'],
        ['Inactive',             'Amber', 'Hidden',  'Registration suspended or on hold'],
        ['Approved Professional','Grey',  'Hidden',  'Approved but not BARB-certified (e.g. SLMC members)'],
    ],
    col_widths=[1.5, 0.9, 1.4, 2.8]
)
add_note_box(doc, 'Changing a therapist\'s status to **Active** automatically makes them **visible in the public directory**. Changing to Inactive or Approved Professional **hides them from the directory**.', kind='note')
add_h3(doc, 'Deleting a therapist')
add_body(doc, 'Each row has a red "Delete" button. Clicking it opens a confirmation modal that requires the admin to **type the therapist\'s full name exactly** before the delete button activates. This is a safety mechanism to prevent accidental deletion.')
add_note_box(doc, '**Warning:** Deleting a therapist is permanent and irreversible. It removes them from the registry and the public directory. Use "Inactive" status instead if you only want to hide them temporarily.', kind='warn')

add_h2(doc, '6.7  Inquiries Dashboard — /admin/inquiries')
add_body(doc, 'This page shows all messages submitted via the Contact form.')
add_h3(doc, 'Features')
for item in [
    '**Filter tabs:** All / New / Read / Replied / Archived — with counts',
    '**Search bar:** Search by name, email, or message content',
    '**Topic filter dropdown:** Filter by inquiry type (Certification, Training, Donation, Collaboration, Complaints, General, Other)',
    '**Status indicator:** New inquiries show a blue dot and bold text; the row has a light blue background',
    '**Expand row:** Click any row to expand and read the full message. Expanding a "New" inquiry automatically marks it as "Read"',
    '**Inline status dropdown:** Change status per inquiry (New → Read → Replied → Archived). The status badge is colour-coded.',
    '**"Reply via Email"** button: Opens the admin\'s default email client pre-filled with the sender\'s address and subject line.',
]:
    add_bullet(doc, item)
add_h3(doc, 'Inquiry statuses')
add_table(doc,
    headers=['Status', 'Colour', 'Meaning'],
    rows=[
        ['New',      'Blue',  'Just received, not yet seen'],
        ['Read',     'Grey',  'Admin has opened and read the inquiry'],
        ['Replied',  'Green', 'Admin has responded to the sender'],
        ['Archived', 'Yellow','No further action needed'],
    ],
    col_widths=[1.1, 1.0, 4.5]
)
add_divider(doc)

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 7 — Database
# ══════════════════════════════════════════════════════════════════════════════

add_h1(doc, '7. Database Architecture')
add_body(doc, 'The database runs on **Supabase (PostgreSQL)**. There are four main tables/views and one OTP helper table.')

for tbl_name, desc, cols, rls in [
    (
        'therapist_applications',
        'Stores every certification application submitted via the website. All fields from the 5-step form map directly to columns here.',
        [
            '`id` (UUID, primary key)',
            '`review_status` — enum: `pending`, `under_review`, `approved`, `rejected`',
            'Personal info fields (first_name, surname, dob, address, phone, email, nic_or_passport)',
            'Criteria boolean flags (current_rbt, current_ibt, expired_rbt, etc.)',
            'Certification numbers (rbt_certification_no, ibt_certification_no, etc.)',
            'Education fields (institution, period_of_education, qualifications)',
            'Work experience fields (work_place_name, work_place_address, designation, etc.)',
            'File path columns (nic_front_file_name, education_file_name, cv_file_name, etc.)',
            'Agreement boolean flags (agree_objectives, agree_ethics, agree_police_clearance, etc.)',
            '`submitted_at`, `updated_at` timestamps',
        ],
        'Anyone (public) can INSERT a new application. Only authenticated admins can SELECT or UPDATE.'
    ),
    (
        'therapists',
        'The regulated registry of approved therapists. A record is created here only when an admin approves an application.',
        [
            '`id` (UUID, primary key)',
            '`application_id` — links back to the source application',
            '`registration_number` — unique BARB registration number (e.g. `RBT-007`)',
            '`full_name` — computed column (first_name + surname, auto-generated by the database)',
            '`status` — enum: `authorized_active`, `unauthorized_inactive`, `approved_non_certified`',
            '`directory_visible` — boolean, controls public directory visibility',
            '`bio`, `profile_image_url` — optional fields editable by admin for directory display',
            '`approved_at`, `last_status_updated_at` timestamps',
        ],
        'Anyone can SELECT therapists where `directory_visible = true`. Only admins can INSERT/UPDATE/DELETE.'
    ),
]:
    add_h2(doc, f'`{tbl_name}`')
    add_body(doc, desc)
    add_body(doc, 'Key columns:')
    for col in cols:
        add_bullet(doc, col)
    add_note_box(doc, f'**RLS Policy:** {rls}', kind='note')

add_h2(doc, '`public_directory` (View)')
add_body(doc, 'A read-only SQL view over the `therapists` table that exposes only safe public fields. The Therapist Directory page queries this view, not the raw `therapists` table.')
add_body(doc, 'Exposes: `id`, `full_name`, `profile_image_url`, `bio`, `registration_number`, `designation`, `work_place_name`, `work_place_address`, `city`, `status`.')

add_h2(doc, '`certification_records`')
add_body(doc, 'Tracks individual certifications and their expiry dates for each therapist. Allows renewals and status changes without altering the core `therapists` row.')

add_h2(doc, '`inquiries`')
add_body(doc, 'Stores all contact form submissions.')
add_body(doc, 'Key columns: `id` (UUID), `name`, `email`, `phone`, `inquiry_type`, `message`, `status` (enum: `new`, `read`, `replied`, `archived`), `created_at`.')

add_h2(doc, '`email_otps`')
add_body(doc, 'Temporary table storing OTP codes for applicant email verification.')
add_body(doc, 'Key columns: `email`, `code`, `expires_at`, `used` (boolean). Old unused codes for an email are deleted before a new one is inserted. Codes expire after 10 minutes.')
add_divider(doc)

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 8 — File Storage
# ══════════════════════════════════════════════════════════════════════════════

add_h1(doc, '8. File Storage')
add_body(doc, 'Uploaded applicant documents are stored in a **Supabase Storage** bucket named **`application-documents`**.')
add_table(doc,
    headers=['Folder', 'Contents'],
    rows=[
        ['`nic/`',            'NIC front and back photos, or passport scans'],
        ['`certifications/`', 'Expired RBT certificates, Expired IBT certificates, Behaviour Analyst documents'],
        ['`education/`',      'Education certificates and transcripts'],
        ['`work/`',           'Work experience letters and references'],
        ['`cv/`',             'Applicant CVs'],
        ['`insurance/`',      'Professional liability / malpractice insurance documents'],
    ],
    col_widths=[1.5, 5.1]
)
add_body(doc, 'File names follow the pattern: `[folder]/[timestamp]_[original filename]`.')
add_body(doc, '**Access:** The bucket is private. Admins can access files via Supabase Dashboard → Storage → `application-documents`, or by generating signed URLs through the Supabase API.')
add_divider(doc)

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 9 — Environment Variables
# ══════════════════════════════════════════════════════════════════════════════

add_h1(doc, '9. Environment Variables & Configuration')
add_body(doc, 'The following environment variables must be set for the website to function. These are stored in `.env.local` locally and must be configured on your hosting platform for production.')
add_table(doc,
    headers=['Variable', 'Where to get it', 'Purpose'],
    rows=[
        ['`NEXT_PUBLIC_SUPABASE_URL`',    'Supabase Dashboard → Project Settings → API → Project URL',         'Connects the frontend to Supabase'],
        ['`NEXT_PUBLIC_SUPABASE_ANON_KEY`','Supabase Dashboard → Project Settings → API → anon public key',   'Public read/write access with RLS enforcement'],
        ['`SUPABASE_SERVICE_ROLE_KEY`',   'Supabase Dashboard → Project Settings → API → service_role key',   'Admin-level server-side access. Never expose this publicly.'],
        ['`RESEND_API_KEY`',              'Resend Dashboard → API Keys',                                       'Sends OTP emails to applicants and inquiry notifications to admins'],
        ['`RESEND_FROM_EMAIL`',           'A verified sender in your Resend account',                          'The "from" address on all outgoing emails (e.g. noreply@barb.lk)'],
        ['`ADMIN_NOTIFICATION_EMAIL`',    'Set manually',                                                      'The email address that receives contact form notifications (e.g. admin@barb.lk)'],
    ],
    col_widths=[2.2, 2.2, 2.2]
)
add_divider(doc)

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 10 — Deployment
# ══════════════════════════════════════════════════════════════════════════════

add_h1(doc, '10. Deployment')
add_body(doc, 'The project ships with a **multi-stage Dockerfile** for containerised deployment. It:')
for i, step in enumerate([
    'Installs only production dependencies',
    'Builds the optimised Next.js application (with `--max-old-space-size=512` for memory-constrained environments)',
    'Produces a minimal Node.js 20 Alpine image running on port 3000',
], 1):
    add_numbered(doc, step, i)

add_h2(doc, 'Running locally (development)')
add_code_block(doc, [
    'npm install',
    'npm run dev',
    '# Visit http://localhost:3000',
])

add_h2(doc, 'Building for production')
add_code_block(doc, [
    'npm run build',
    'npm run start',
])

add_h2(doc, 'Building with Docker')
add_code_block(doc, [
    'docker build -t barb-website .',
    'docker run -p 3000:3000 --env-file .env.local barb-website',
])

add_h2(doc, 'Recommended hosting options')
add_table(doc,
    headers=['Option', 'Notes'],
    rows=[
        ['Vercel',          'Easiest — connect the GitHub repo, add environment variables in the Vercel dashboard, and it deploys automatically on every push. Zero configuration needed.'],
        ['VPS (Docker)',     'Use the Dockerfile. Recommended if hosting on a Sri Lankan cloud provider or private server. Requires Nginx reverse proxy and SSL (Let\'s Encrypt).'],
        ['Fly.io / Railway', 'Supports Docker deployments with easy environment variable management.'],
    ],
    col_widths=[1.6, 5.0]
)
add_divider(doc)

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 11 — Credentials
# ══════════════════════════════════════════════════════════════════════════════

add_h1(doc, '11. Credentials & Access Handover')
add_body(doc, 'The following access must be transferred to the BARB team. Semantix Labs will provide the actual credentials/keys separately and securely.')
add_table(doc,
    headers=['Service', 'Access Type', 'Notes'],
    rows=[
        ['**Supabase project**',      'Project Owner / Admin',     'Transfer project ownership in Supabase Dashboard → Settings → Team'],
        ['**Resend account**',        'API Key access',            'Create a new API key scoped to the production domain and revoke the development one'],
        ['**Domain / DNS**',          'Registrar access',          'Configure A/CNAME records to point to the hosting server'],
        ['**Hosting / Server**',      'SSH or platform login',     'Depends on chosen deployment method'],
        ['**Admin user account**',    'Email + password (Supabase Auth)', 'Create at: Supabase Dashboard → Authentication → Users → Add user'],
        ['**Source code**',           'GitHub repository access',  'Transfer repo ownership or add BARB team members as collaborators'],
    ],
    col_widths=[1.7, 1.8, 3.1]
)
add_divider(doc)

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 12 — Maintenance
# ══════════════════════════════════════════════════════════════════════════════

add_h1(doc, '12. Ongoing Maintenance Notes')

tasks = [
    ('Adding a new admin user', [
        'Go to Supabase Dashboard → Authentication → Users',
        'Click "Add user" and enter the new admin\'s email and a temporary password',
        'Share the credentials securely with the new admin',
        'They log in at `/admin/login`',
    ], None),
    ('Resetting an admin password', None, 'Supabase can send a recovery email from the Authentication dashboard. The admin clicks the link in the email, which takes them to `/admin/reset-password` where they set a new password.'),
    ('Managing therapist directory visibility', [
        'Use the **Therapist Registry** page (`/admin/therapists`) to change any therapist\'s status',
        'Set to **Active** → therapist appears in the public directory',
        'Set to **Inactive** or **Approved Professional** → therapist is hidden from the directory',
    ], None),
    ('Updating a therapist\'s bio or profile image', [
        'Go to Supabase Dashboard → Table Editor → `therapists`',
        'Find the therapist row and edit the `bio` or `profile_image_url` columns',
        'For profile images: upload the image to Supabase Storage, copy the URL, and paste it into `profile_image_url`',
    ], None),
    ('Permanently removing a therapist', None, 'Use the Delete button on the **Therapist Registry** page. You must type the therapist\'s full name in the confirmation modal. This is permanent — use "Inactive" status if you only want to hide them temporarily.'),
    ('Changing the admin notification email', None, 'Update the `ADMIN_NOTIFICATION_EMAIL` environment variable in your hosting platform settings and redeploy.'),
    ('Changing the "from" email address', [
        'Add and verify the new domain/address in your Resend account',
        'Update `RESEND_FROM_EMAIL` in your environment variables and redeploy',
    ], None),
    ('What if the OTP email is not received?', [
        'Check the Resend dashboard for delivery errors',
        'Confirm `RESEND_API_KEY` and `RESEND_FROM_EMAIL` are correctly set',
        'Ensure the sender domain is verified in Resend',
        'Check the `email_otps` table in Supabase to confirm the code was generated',
    ], None),
    ('Database backups', None, 'Supabase automatically takes daily backups on paid plans. You can also export the database manually via Supabase Dashboard → Database → Backups.'),
]

for title, bullets, prose in tasks:
    add_h2(doc, title)
    if prose:
        add_body(doc, prose)
    if bullets:
        for item in bullets:
            add_bullet(doc, item)

add_divider(doc)

# ── Footer paragraph ───────────────────────────────────────────────────────────
p = doc.add_paragraph()
set_para_spacing(p, before=8, after=0)
r = p.add_run('This document was prepared by Semantix Labs as part of the project handover to BARB SL. For questions about this document or the codebase, contact Semantix Labs at semantixlabs@gmail.com.')
r.font.size = Pt(9)
r.font.color.rgb = GRAY_TXT
r.italic = True

# ── Save ───────────────────────────────────────────────────────────────────────
out = '/Users/shobians./Desktop/Semantix/BARB/Website/BTRB-V1/_docs/BARB_SL_Website_Handover_Document.docx'
doc.save(out)
print(f'Saved → {out}')
